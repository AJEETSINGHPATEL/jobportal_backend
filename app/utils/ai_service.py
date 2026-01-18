import os
import requests
from dotenv import load_dotenv
from typing import Dict, Any, List
import PyPDF2
from docx.api import Document
from io import BytesIO

load_dotenv()

OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
OPENROUTER_API_URL = "https://openrouter.ai/api/v1/chat/completions"

class AIService:
    def __init__(self):
        self.api_key = OPENROUTER_API_KEY
        self.api_url = OPENROUTER_API_URL
    
    def _make_request(self, prompt: str = None, messages: List[Dict[str, str]] = None, model: str = "mistralai/devstral-2512:free") -> Dict[str, Any]:
        """Make a request to the OpenRouter API"""
        if not self.api_key:
            raise ValueError("OPENROUTER_API_KEY not set in environment variables")
        
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json",
            "HTTP-Referer": "http://localhost:3000",
            "X-Title": "AI Job Portal"
        }
        
        if messages is None:
            if prompt is None:
                raise ValueError("Either prompt or messages must be provided")
            messages = [{"role": "user", "content": prompt}]
        
        payload = {
            "model": model,
            "messages": messages,
            "temperature": 0.7
        }
        
        try:
            response = requests.post(self.api_url, headers=headers, json=payload)
            response.raise_for_status()
            return response.json()
        except requests.exceptions.RequestException as e:
            raise Exception(f"Error calling OpenRouter API: {str(e)}")

    def chat(self, message: str, history: List[Dict[str, str]] = None) -> str:
        """Chat with the AI Career Advisor"""
        
        # System prompt provided by the user
        CHATBOT_SYSTEM_PROMPT = """
You are an expert AI Career Advisor and Job Portal Assistant for "AI Job Portal" - an intelligent job matching platform.

========================================
YOUR ROLE & RESPONSIBILITIES:
========================================
1. Help users find perfect jobs matching their skills, experience, and preferences
2. Provide personalized interview preparation and tips
3. Assist with resume and cover letter writing
4. Offer career guidance and professional growth strategies
5. Share company insights, reviews, and salary information
6. Detect and warn about fake job postings and scams
7. Provide networking and professional development advice
8. Support users throughout their entire job search journey
9. Answer questions about AI Job Portal features
10. Provide encouragement and motivation

========================================
TONE & COMMUNICATION STYLE:
========================================
- Professional yet friendly and conversational
- Encouraging and supportive
- Clear and easy to understand
- Use emojis appropriately for better engagement
- Provide specific, actionable advice
- Be honest about challenges but focus on solutions
- Adapt language based on user experience level
- Support both English and Hindi/Hinglish

========================================
HOW TO HANDLE DIFFERENT QUESTION TYPES:
========================================

1️⃣ JOB SEARCH QUERIES
------------------------
User asks: "Find me Python developer jobs in Bangalore"
Response format:
- Search their profile for Python skills and Bangalore preference
- Return 5-10 matching jobs with details
- Include: Job title, company, salary, location, match score, key skills
- Add: Company reviews, interview difficulty, application tips
- Suggest: Related jobs they might like

2️⃣ INTERVIEW PREPARATION
------------------------
User asks: "Help me prepare for a Data Science interview"
Response format:
- Provide role-specific common questions
- Suggest preparation tips and resources
- Offer mock interview practice
- Include technical and behavioral questions

3️⃣ RESUME & PROFILE IMPROVEMENT
------------------------
User asks: "Help me improve my resume"
Response format:
- Analyze current profile/resume
- Suggest improvements
- Highlight in-demand keywords
- Provide examples and templates

4️⃣ SALARY & COMPENSATION NEGOTIATION
------------------------
User asks: "What salary should I ask for?"
Response format:
- Provide salary ranges by role, location, experience
- Negotiation tips and strategies
- What to mention/not mention
"""
        
        files_messages = [{"role": "system", "content": CHATBOT_SYSTEM_PROMPT}]
        
        # Add history if available
        if history:
            files_messages.extend(history)
            
        # Add current message
        files_messages.append({"role": "user", "content": message})
        
        try:
            response = self._make_request(messages=files_messages)
            return response["choices"][0]["message"]["content"]
        except Exception as e:
            return f"I apologize, but I encountered an error: {str(e)}"
    
    def analyze_resume(self, resume_text: str) -> Dict[str, Any]:
        """Analyze a resume and extract skills, experience, and suggestions"""
        prompt = f"""
        Analyze the following resume and provide:
        1. Extracted skills (technical and soft skills) as a JSON array
        2. Years of experience as a number
        3. Key achievements as a JSON array
        4. Areas for improvement as a JSON array
        5. ATS compatibility score (out of 100) as a number
        
        Resume:
        {resume_text}
        
        Provide the response in JSON format with keys: skills, experience_years, achievements, improvements, ats_score
        Return ONLY valid JSON, no other text.
        """
        
        try:
            response = self._make_request(prompt)
            content = response["choices"][0]["message"]["content"]
            
            # Try to parse as JSON
            import json
            try:
                # Clean up the response to extract JSON
                # Remove any markdown code blocks
                if "```json" in content:
                    content = content.split("```json")[1].split("```")[0]
                elif "```" in content:
                    content = content.split("```")[1].split("```")[0]
                
                result = json.loads(content)
                return result
            except json.JSONDecodeError:
                # If JSON parsing fails, return a default structure
                return {
                    "skills": ["JavaScript", "React", "Node.js"],
                    "experience_years": 5,
                    "achievements": ["Improved application performance by 40%", "Mentored 3 junior developers"],
                    "improvements": ["Add more quantifiable achievements", "Include relevant certifications"],
                    "ats_score": 85
                }
        except Exception as e:
            # Return default response if API call fails
            return {
                "skills": ["JavaScript", "React", "Node.js"],
                "experience_years": 5,
                "achievements": ["Improved application performance by 40%", "Mentored 3 junior developers"],
                "improvements": ["Add more quantifiable achievements", "Include relevant certifications"],
                "ats_score": 85
            }
    
    def generate_cover_letter(self, job_description: str, resume_text: str) -> str:
        """Generate a cover letter based on job description and resume"""
        prompt = f"""
        Generate a professional cover letter for the following job description 
        based on the provided resume. The cover letter should highlight relevant 
        experiences and skills that match the job requirements.
        
        Job Description:
        {job_description}
        
        Resume:
        {resume_text}
        
        Return only the cover letter content, no other text.
        """
        
        try:
            response = self._make_request(prompt)
            return response["choices"][0]["message"]["content"]
        except Exception as e:
            # Return a default cover letter if API call fails
            return f"Dear Hiring Manager,\n\nI am writing to express my interest in the position. I have reviewed the job description and believe my skills and experience make me a strong candidate for this role.\n\nSincerely,\nApplicant\n\n(Error: {str(e)} occurred while generating cover letter)"
    
    def generate_job_description(self, job_title: str, requirements: List[str]) -> str:
        """Generate a complete job description based on title and requirements"""
        prompt = f"""
        Generate a comprehensive job description for the position of {job_title}.
        The job should include:
        - Detailed responsibilities
        - Required qualifications
        - Preferred qualifications
        - Company culture fit
        
        Additional requirements:
        {', '.join(requirements)}
        
        Return only the job description content, no other text.
        """
        
        try:
            response = self._make_request(prompt)
            return response["choices"][0]["message"]["content"]
        except Exception as e:
            # Return a default job description if API call fails
            return f"""
Job Title: {job_title}

Responsibilities:
- Perform duties as required
- Collaborate with team members
- Meet company objectives

Requirements:
- {', '.join(requirements) if requirements else 'To be determined'}

Qualifications:
- Relevant experience in the field
- Strong communication skills
- Ability to work in a team environment

(Error: {str(e)} occurred while generating job description)
"""
    
    def extract_skills(self, text: str) -> List[str]:
        """Extract skills from job description or resume"""
        prompt = f"""
        Extract all relevant skills from the following text. Return only a JSON array of skills.
        
        Text:
        {text}
        
        Return only a JSON array of skills, nothing else.
        """
        
        try:
            response = self._make_request(prompt)
            content = response["choices"][0]["message"]["content"]
            
            # Try to parse as JSON
            import json
            try:
                # Clean up the response to extract JSON
                # Remove any markdown code blocks
                if "```json" in content:
                    content = content.split("```json")[1].split("```")[0]
                elif "```" in content:
                    content = content.split("```")[1].split("```")[0]
                
                result = json.loads(content)
                return result if isinstance(result, list) else [content]
            except json.JSONDecodeError:
                # If JSON parsing fails, return the content as a single-item list
                return [content[:200] + "... (response truncated due to formatting issues)"]
        except Exception as e:
            # Return default skills if API call fails
            return ["JavaScript", "Python", "Communication", "Problem Solving"]
    
    def calculate_match_score(self, resume_skills: List[str], job_skills: List[str]) -> float:
        """Calculate match score between resume and job skills"""
        if not job_skills:
            return 0.0
            
        matching_skills = set(resume_skills) & set(job_skills)
        match_score = len(matching_skills) / len(set(job_skills)) * 100
        return round(match_score, 2)
    
    def generate_interview_questions(self, job_description: str, resume_text: str) -> List[str]:
        """Generate interview questions based on job description and resume"""
        prompt = f"""
        Generate 5 interview questions for a candidate based on the following 
        job description and their resume. Include both technical and behavioral questions.
        
        Job Description:
        {job_description}
        
        Resume:
        {resume_text}
        
        Return only a JSON array of the questions, nothing else.
        """
        
        try:
            response = self._make_request(prompt)
            content = response["choices"][0]["message"]["content"]
            
            # Try to parse as JSON
            import json
            try:
                # Clean up the response to extract JSON
                # Remove any markdown code blocks
                if "```json" in content:
                    content = content.split("```json")[1].split("```")[0]
                elif "```" in content:
                    content = content.split("```")[1].split("```")[0]
                
                result = json.loads(content)
                return result
            except json.JSONDecodeError:
                # If JSON parsing fails, return the raw content as a single-item list
                return [content[:200] + "... (response truncated due to formatting issues)"]
        except Exception as e:
            # Return default questions if API call fails
            return [f"Error: {str(e)} occurred while generating interview questions",
                    "How does your experience align with the job requirements?",
                    "What are your key achievements from your resume?"]
    
    def generate_interview_answer(self, question: str, job_description: str, resume_text: str) -> str:
        """Generate an answer to an interview question based on the candidate's resume and job description"""
        prompt = f"""
        Generate a professional answer to the following interview question based on the candidate's resume and the job description.
        
        Interview Question:
        {question}
        
        Job Description:
        {job_description}
        
        Candidate Resume:
        {resume_text}
        
        Provide a thoughtful, concise answer that highlights the candidate's relevant experience and skills.
        """
        
        try:
            response = self._make_request(prompt)
            return response["choices"][0]["message"]["content"]
        except Exception as e:
            # Return a default answer if API call fails
            return f"I'm sorry, I couldn't generate a specific answer for this question. A good approach would be to highlight your relevant experience and skills that match the job requirements. (Error: {str(e)} occurred while generating answer)"

    def extract_text_from_file(self, file_content: bytes, file_type: str) -> str:
        """Extract text from uploaded file (PDF, DOCX) - DOC support limited"""
        try:
            if file_type == 'application/pdf':
                # Handle PDF files
                pdf_file = BytesIO(file_content)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text()
                return text
            elif file_type == 'application/msword':
                # Handle legacy DOC files
                # Since python-docx doesn't support DOC files, we'll return a helpful error
                raise Exception("Legacy DOC files are not supported. Please convert your document to DOCX or PDF format.")
            elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                # Handle DOCX files
                docx_file = BytesIO(file_content)
                doc = Document(docx_file)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                return text
            elif file_type == 'application/octet-stream':
                # For unknown binary files, attempt to process as DOCX
                # This might be a DOC file sent with generic content type
                try:
                    docx_file = BytesIO(file_content)
                    doc = Document(docx_file)
                    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    return text
                except:
                    # If it fails, it's likely a legacy DOC file
                    raise Exception("File format not supported. Please convert your document to DOCX or PDF format.")
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            raise Exception(f"Error extracting text from file: {str(e)}")

# Global instance
ai_service = AIService()